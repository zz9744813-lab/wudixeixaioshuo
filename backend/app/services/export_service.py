"""
Export Service - 小说导出服务
支持多种格式导出
"""

import json
import logging
import re
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import List, Optional

from sqlalchemy.orm import Session

from app.models.chapter import Chapter, ChapterStatus
from app.models.project import Project

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    """导出格式"""
    MARKDOWN = "md"
    TXT = "txt"
    DOCX = "docx"
    EPUB = "epub"
    PDF = "pdf"
    JSON = "json"


class ExportService:
    """
    小说导出服务

    支持格式：
    - Markdown: 带格式标记的文本
    - TXT: 纯文本
    - DOCX: Word 文档
    - EPUB: 电子书格式
    - PDF: 打印格式
    - JSON: 结构化数据
    """

    def __init__(self, db: Session):
        self.db = db
        self.export_dir = Path("exports")
        self.export_dir.mkdir(exist_ok=True)

    def export_project(
        self,
        project_id: int,
        format: ExportFormat,
        include_outline: bool = True,
        include_metadata: bool = True,
        chapter_filter: Optional[str] = "completed"  # all, completed, reviewed
    ) -> dict:
        """
        导出项目

        Args:
            project_id: 项目ID
            format: 导出格式
            include_outline: 是否包含大纲
            include_metadata: 是否包含元数据
            chapter_filter: 章节过滤条件
        """
        project = self.db.query(Project).filter(
            Project.id == project_id
        ).first()

        if not project:
            return {"error": "项目不存在"}

        # 获取章节
        query = self.db.query(Chapter).filter(
            Chapter.project_id == project_id
        )

        if chapter_filter == "completed":
            query = query.filter(Chapter.status == ChapterStatus.COMPLETED)
        elif chapter_filter == "reviewed":
            query = query.filter(Chapter.status.in_([
                ChapterStatus.REVIEW,
                ChapterStatus.COMPLETED
            ]))

        chapters = query.order_by(Chapter.order_num.asc()).all()

        if not chapters:
            return {"error": "没有可导出的章节"}

        # 根据格式导出
        exporters = {
            ExportFormat.MARKDOWN: self._export_markdown,
            ExportFormat.TXT: self._export_txt,
            ExportFormat.DOCX: self._export_docx,
            ExportFormat.EPUB: self._export_epub,
            ExportFormat.PDF: self._export_pdf,
            ExportFormat.JSON: self._export_json,
        }

        exporter = exporters.get(format)
        if not exporter:
            return {"error": f"不支持的格式: {format}"}

        try:
            result = exporter(project, chapters, include_outline, include_metadata)
            return result
        except Exception as e:
            logger.error(f"导出失败: {e}")
            return {"error": f"导出失败: {str(e)}"}

    def _export_markdown(
        self,
        project: Project,
        chapters: List[Chapter],
        include_outline: bool,
        include_metadata: bool
    ) -> dict:
        """导出为 Markdown"""
        lines = []

        # 标题
        lines.append(f"# {project.name}")
        lines.append("")

        # 元数据
        if include_metadata:
            lines.append("---")
            lines.append(f"作者: {project.config.get('author', 'AI生成')}")
            lines.append(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            lines.append(f"总章节数: {len(chapters)}")
            lines.append(f"总字数: {sum(c.word_count or 0 for c in chapters)}")
            lines.append("---")
            lines.append("")

        # 大纲
        if include_outline and project.outline:
            lines.append("## 大纲")
            lines.append(project.outline)
            lines.append("")
            lines.append("---")
            lines.append("")

        # 目录
        lines.append("## 目录")
        for i, ch in enumerate(chapters, 1):
            lines.append(f"{i}. {ch.title}")
        lines.append("")
        lines.append("---")
        lines.append("")

        # 章节内容
        for ch in chapters:
            lines.append(f"## {ch.title}")
            lines.append("")
            if ch.content:
                lines.append(ch.content)
            else:
                lines.append("*章节内容待生成*")
            lines.append("")
            lines.append("---")
            lines.append("")

        content = "\n".join(lines)

        # 保存文件
        filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        filepath = self.export_dir / filename
        filepath.write_text(content, encoding='utf-8')

        return {
            "format": "markdown",
            "filename": filename,
            "filepath": str(filepath),
            "chapter_count": len(chapters),
            "word_count": sum(c.word_count or 0 for c in chapters),
        }

    def _export_txt(
        self,
        project: Project,
        chapters: List[Chapter],
        include_outline: bool,
        include_metadata: bool
    ) -> dict:
        """导出为纯文本"""
        lines = []

        lines.append(project.name)
        lines.append("=" * len(project.name))
        lines.append("")

        if include_metadata:
            lines.append(f"作者: {project.config.get('author', 'AI生成')}")
            lines.append(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            lines.append("")

        if include_outline and project.outline:
            lines.append("【大纲】")
            lines.append(project.outline)
            lines.append("")

        for ch in chapters:
            lines.append(f"\n{'='*40}")
            lines.append(f"第 {ch.order_num} 章")
            lines.append(ch.title)
            lines.append(f"{'='*40}\n")

            if ch.content:
                # 移除 Markdown 标记
                text = re.sub(r'#+ ', '', ch.content)
                text = re.sub(r'\*\*|__', '', text)
                text = re.sub(r'`', '', text)
                lines.append(text)
            else:
                lines.append("【章节内容待生成】")

            lines.append("")

        content = "\n".join(lines)

        filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = self.export_dir / filename
        filepath.write_text(content, encoding='utf-8')

        return {
            "format": "txt",
            "filename": filename,
            "filepath": str(filepath),
            "chapter_count": len(chapters),
            "word_count": sum(c.word_count or 0 for c in chapters),
        }

    def _export_json(
        self,
        project: Project,
        chapters: List[Chapter],
        include_outline: bool,
        include_metadata: bool
    ) -> dict:
        """导出为 JSON"""
        data = {
            "project": {
                "id": project.id,
                "name": project.name,
                "description": project.description,
                "genre": project.genre,
                "target_words": project.target_words,
                "config": project.config if include_metadata else None,
            },
            "export_info": {
                "format": "json",
                "exported_at": datetime.now().isoformat(),
                "chapter_count": len(chapters),
                "total_word_count": sum(c.word_count or 0 for c in chapters),
            },
            "outline": project.outline if include_outline else None,
            "chapters": [
                {
                    "id": ch.id,
                    "order_num": ch.order_num,
                    "title": ch.title,
                    "content": ch.content,
                    "word_count": ch.word_count,
                    "status": ch.status.value if ch.status else None,
                    "metadata": ch.metadata,
                }
                for ch in chapters
            ]
        }

        content = json.dumps(data, ensure_ascii=False, indent=2)

        filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        filepath = self.export_dir / filename
        filepath.write_text(content, encoding='utf-8')

        return {
            "format": "json",
            "filename": filename,
            "filepath": str(filepath),
            "chapter_count": len(chapters),
            "word_count": sum(c.word_count or 0 for c in chapters),
        }

    def _export_docx(
        self,
        project: Project,
        chapters: List[Chapter],
        include_outline: bool,
        include_metadata: bool
    ) -> dict:
        """导出为 Word 文档（真实实现）"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx 未安装，回退到 Markdown 导出")
            result = self._export_markdown(project, chapters, include_outline, include_metadata)
            result["format"] = "docx"
            result["filename"] = result["filename"].replace(".md", ".docx")
            result["note"] = "python-docx 未安装，请使用 pip install python-docx"
            return result

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

        # 标题
        title = doc.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元数据
        if include_metadata:
            doc.add_paragraph()
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"作者: {project.config.get('author', 'AI生成') if project.config else 'AI生成'}").font.size = Pt(10)
            doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", style='Normal')
            doc.add_paragraph(f"总章节数: {len(chapters)}", style='Normal')
            doc.add_paragraph(f"总字数: {sum(c.word_count or 0 for c in chapters)}", style='Normal')
            doc.add_paragraph()

        # 大纲
        if include_outline and project.outline:
            doc.add_heading('大纲', 1)
            doc.add_paragraph(project.outline)
            doc.add_page_break()

        # 目录
        doc.add_heading('目录', 1)
        for i, ch in enumerate(chapters, 1):
            doc.add_paragraph(f"第{i}章 {ch.title}", style='List Number')
        doc.add_page_break()

        # 章节内容
        for ch in chapters:
            doc.add_heading(ch.title, 1)

            if ch.content:
                # 简单处理 Markdown 标记
                paragraphs = ch.content.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue

                    # 处理标题
                    if para_text.startswith('# '):
                        doc.add_heading(para_text[2:], 2)
                    elif para_text.startswith('## '):
                        doc.add_heading(para_text[3:], 3)
                    elif para_text.startswith('### '):
                        doc.add_heading(para_text[4:], 4)
                    else:
                        # 普通段落，移除 Markdown 标记
                        para_text = re.sub(r'\*\*|__', '', para_text)  # 粗体
                        para_text = re.sub(r'\*|_', '', para_text)     # 斜体
                        para_text = re.sub(r'`', '', para_text)        # 代码
                        doc.add_paragraph(para_text)
            else:
                doc.add_paragraph('（章节内容待生成）', style='Intense Quote')

            doc.add_paragraph()  # 章节间隔

        # 保存文件
        filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.export_dir / filename
        doc.save(str(filepath))

        return {
            "format": "docx",
            "filename": filename,
            "filepath": str(filepath),
            "chapter_count": len(chapters),
            "word_count": sum(c.word_count or 0 for c in chapters),
        }

    def _export_epub(
        self,
        project: Project,
        chapters: List[Chapter],
        include_outline: bool,
        include_metadata: bool
    ) -> dict:
        """导出为 EPUB（真实实现）"""
        try:
            from ebooklib import epub
        except ImportError:
            logger.error("ebooklib 未安装，回退到 Markdown 导出")
            result = self._export_markdown(project, chapters, include_outline, include_metadata)
            result["format"] = "epub"
            result["filename"] = result["filename"].replace(".md", ".epub")
            result["note"] = "ebooklib 未安装，请使用 pip install ebooklib"
            return result

        # 创建 EPUB 书籍
        book = epub.EpubBook()

        # 设置元数据
        book.set_identifier(f'novel-{project.id}-{datetime.now().timestamp()}')
        book.set_title(project.name)
        book.set_language('zh-CN')

        author = project.config.get('author', 'AI生成') if project.config else 'AI生成'
        book.add_author(author)

        # 简介
        if project.description:
            book.add_metadata('DC', 'description', project.description)

        # 创建章节
        epub_chapters = []
        toc = []

        # 简介/大纲章节
        if include_outline and project.outline:
            intro = epub.EpubHtml(title='简介', file_name='intro.xhtml', lang='zh-CN')
            intro.content = f'<h1>简介</h1><p>{project.outline}</p>'
            book.add_item(intro)
            epub_chapters.append(intro)
            toc.append(intro)

        # 添加各章节
        for i, ch in enumerate(chapters, 1):
            chapter_file = f'chapter_{i:04d}.xhtml'
            epub_ch = epub.EpubHtml(
                title=ch.title,
                file_name=chapter_file,
                lang='zh-CN'
            )

            # 构建章节内容
            content_parts = [f'<h1>{ch.title}</h1>']

            if ch.content:
                # 简单 Markdown 转 HTML
                html_content = self._markdown_to_html(ch.content)
                content_parts.append(html_content)
            else:
                content_parts.append('<p><em>章节内容待生成</em></p>')

            epub_ch.content = '\n'.join(content_parts)
            book.add_item(epub_ch)
            epub_chapters.append(epub_ch)
            toc.append(epub_ch)

        # 添加 TOC
        book.toc = toc

        # 添加 NCX 和导航
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # 定义 spine（阅读顺序）
        book.spine = ['nav'] + epub_chapters

        # 添加默认 CSS
        style = '''
        body { font-family: serif; line-height: 1.6; }
        h1 { font-size: 1.8em; margin: 1em 0; }
        h2 { font-size: 1.5em; margin: 0.8em 0; }
        p { margin: 0.5em 0; text-indent: 2em; }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)

        # 保存文件
        filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.epub"
        filepath = self.export_dir / filename
        epub.write_epub(str(filepath), book, {})

        return {
            "format": "epub",
            "filename": filename,
            "filepath": str(filepath),
            "chapter_count": len(chapters),
            "word_count": sum(c.word_count or 0 for c in chapters),
        }

    def _markdown_to_html(self, markdown_text: str) -> str:
        """简单 Markdown 转 HTML"""
        import html

        # 转义 HTML 特殊字符
        text = html.escape(markdown_text)

        # 处理标题
        text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
        text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)

        # 处理粗体和斜体
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)

        # 处理段落
        paragraphs = text.split('\n\n')
        html_parts = []
        for para in paragraphs:
            para = para.strip()
            if para and not para.startswith('<'):
                html_parts.append(f'<p>{para}</p>')
            else:
                html_parts.append(para)

        return '\n'.join(html_parts)

    def _export_pdf(
        self,
        project: Project,
        chapters: List[Chapter],
        include_outline: bool,
        include_metadata: bool
    ) -> dict:
        """导出为 PDF（使用 reportlab 实现）"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        except ImportError:
            logger.error("reportlab 未安装，回退到 Markdown 导出")
            result = self._export_markdown(project, chapters, include_outline, include_metadata)
            result["format"] = "pdf"
            result["filename"] = result["filename"].replace(".md", ".pdf")
            result["note"] = "reportlab 未安装，请使用 pip install reportlab"
            return result

        # 注册中文字体（尝试常用字体）
        chinese_font = 'Helvetica'
        try:
            # 尝试注册系统中可能存在的中文字体
            font_paths = [
                'C:/Windows/Fonts/simhei.ttf',  # Windows 黑体
                'C:/Windows/Fonts/simsun.ttc',  # Windows 宋体
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # Linux
                '/System/Library/Fonts/PingFang.ttc',  # macOS
            ]
            for font_path in font_paths:
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    chinese_font = 'ChineseFont'
                    break
                except Exception:
                    continue
        except Exception as e:
            logger.warning(f"无法注册中文字体: {e}")

        # 创建 PDF 文档
        filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = self.export_dir / filename

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        # 创建样式
        styles = getSampleStyleSheet()

        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER
        )

        # 章节标题样式
        chapter_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=18,
            spaceAfter=12
        )

        # 正文样式
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontName=chinese_font,
            fontSize=12,
            leading=20,
            alignment=TA_JUSTIFY
        )

        # 构建内容
        story = []

        # 标题
        story.append(Paragraph(project.name, title_style))
        story.append(Spacer(1, 20))

        # 元数据
        if include_metadata:
            author = project.config.get('author', 'AI生成') if project.config else 'AI生成'
            story.append(Paragraph(f"作者: {author}", body_style))
            story.append(Paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
            story.append(Paragraph(f"总章节数: {len(chapters)}", body_style))
            story.append(Paragraph(f"总字数: {sum(c.word_count or 0 for c in chapters)}", body_style))
            story.append(Spacer(1, 20))

        # 大纲
        if include_outline and project.outline:
            story.append(Paragraph('大纲', chapter_style))
            story.append(Paragraph(project.outline, body_style))
            story.append(PageBreak())

        # 章节内容
        for ch in chapters:
            story.append(Paragraph(ch.title, chapter_style))
            story.append(Spacer(1, 10))

            if ch.content:
                # 处理 Markdown 并分段
                paragraphs = self._prepare_pdf_paragraphs(ch.content)
                for para in paragraphs:
                    story.append(Paragraph(para, body_style))
                    story.append(Spacer(1, 6))
            else:
                story.append(Paragraph('（章节内容待生成）', body_style))

            story.append(Spacer(1, 20))

        # 生成 PDF
        doc.build(story)

        return {
            "format": "pdf",
            "filename": filename,
            "filepath": str(filepath),
            "chapter_count": len(chapters),
            "word_count": sum(c.word_count or 0 for c in chapters),
        }

    def _prepare_pdf_paragraphs(self, content: str) -> List[str]:
        """准备 PDF 段落文本"""
        import html

        paragraphs = []
        for para in content.split('\n\n'):
            para = para.strip()
            if not para:
                continue

            # 简单 Markdown 处理
            para = re.sub(r'^#+ ', '', para)  # 移除标题标记
            para = re.sub(r'\*\*|__', '', para)  # 移除粗体
            para = re.sub(r'\*|_', '', para)     # 移除斜体
            para = re.sub(r'`', '', para)        # 移除代码

            # 转义 HTML
            para = html.escape(para)
            paragraphs.append(para)

        return paragraphs

    def get_export_history(
        self,
        project_id: Optional[int] = None,
        limit: int = 20
    ) -> List[dict]:
        """获取导出历史"""
        exports = []

        for filepath in sorted(self.export_dir.glob("*"), reverse=True)[:limit]:
            if filepath.is_file():
                stat = filepath.stat()
                exports.append({
                    "filename": filepath.name,
                    "filepath": str(filepath),
                    "size": stat.st_size,
                    "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "format": filepath.suffix.lstrip("."),
                })

        return exports

    def delete_export(self, filename: str) -> bool:
        """删除导出文件"""
        filepath = self.export_dir / filename
        if filepath.exists():
            filepath.unlink()
            return True
        return False

    def get_word_count_stats(self, project_id: int) -> dict:
        """获取字数统计"""
        chapters = self.db.query(Chapter).filter(
            Chapter.project_id == project_id
        ).all()

        total = sum(c.word_count or 0 for c in chapters)
        completed = sum(
            c.word_count or 0 for c in chapters
            if c.status == ChapterStatus.COMPLETED
        )

        return {
            "total_chapters": len(chapters),
            "total_word_count": total,
            "completed_word_count": completed,
            "average_per_chapter": round(total / len(chapters), 0) if chapters else 0,
            "by_status": {
                "completed": sum(1 for c in chapters if c.status == ChapterStatus.COMPLETED),
                "writing": sum(1 for c in chapters if c.status == ChapterStatus.WRITING),
                "pending": sum(1 for c in chapters if c.status == ChapterStatus.PENDING),
            }
        }
